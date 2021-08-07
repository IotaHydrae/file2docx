# coding:utf-8

from docx import Document
import os
import sys
from PyQt5.QtWidgets import *
from PyQt5 import QtGui
from GUI import Ui_Dialog
import time

class UserWidgets(QMainWindow, Ui_Dialog):
    def __init__(self, *args, **kwargs):
        super(UserWidgets, self).__init__(*args, **kwargs)
        self.pwd = os.getcwd()
        self.files_list = []
        self.document = Document()
        self.init_menu()


        self.setupUi(self)
        self.config()
        self.set_events()
        self.changes()

    def init_menu(self):
        bar = self.menuBar()
        stop = bar.addMenu('文件')
        settings = bar.addMenu('设置')

        stop.addAction('退出',exit)
        filetype = QAction('文件类型',self)
        filetype.setShortcut("Ctrl + F")
        filetype.triggered.connect(self.get_filetype)


        settings.addAction(filetype)

    def changes(self):
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap("images/dict.jpg"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.project.setIcon(icon)
        self.output.setIcon(icon)

    def config(self):
        self.progressBar.setVisible(False)
        if os.path.exists('cache/project.path'):
            with open('cache/project.path', 'r') as file:

                self.pro_dict.insert(file.read())
                file.close()

        if os.path.exists('cache/output.path'):
            with open('cache/output.path', 'r') as file:

                self.output_dict.insert(file.read())
                file.close()

    def set_events(self):
        self.project.clicked.connect(self.choose_pro_dir)
        self.output.clicked.connect(self.choose_output_dir)
        self.reset.clicked.connect(self.dialog_reset)
        self.generate.clicked.connect(self.generate_docx)

    def choose_pro_dir(self):
        self.progressBar.setVisible(False)
        self.pro_dict.clear()
        pro_path = QFileDialog.getExistingDirectory(self, '请选择项目路径', './')
        self.pro_dict.insert(pro_path)
        self.write_to_file('project', 'path', pro_path)

    def choose_output_dir(self):
        self.progressBar.setVisible(False)
        self.output_dict.clear()
        output_path = QFileDialog.getExistingDirectory(self, '请选择生成位置', './')
        self.output_dict.insert(output_path)
        self.write_to_file('output','path',output_path)

    def get_filetype(self):
        status = os.path.exists('cache/file.tp')
        if status:
            str = self.read_file('cache/file.tp')
            filetype = QInputDialog.getText(self, "文件类型","请输入文件类型(.*)以分号隔开", QLineEdit.Normal,str)
        else:
            filetype = QInputDialog.getText(self, "文件类型", "请输入文件类型(.*)以分号隔开", QLineEdit.Normal )
        if filetype[1]:
            self.write_to_file('file', 'tp', filetype[0])

    def dialog_reset(self):
        self.pro_dict.clear()
        self.output_dict.clear()
        self.progressBar.setVisible(False)

    def visitDir(self, path):
        list_all = os.listdir(path)

        for item in list_all:
            name = os.path.join(path, item)
            if not os.path.isfile(name):
                self.visitDir(name)
            else:
                if name != os.getcwd() +'\\'+ os.path.basename(__file__):
                    # print(name)
                    self.files_list.append(name)
        return self.files_list


    def read_file(self, path):
        str = ''

        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as file:
                str = file.read()
                file.close()

        return str
    def write_to_file(self,filename,type,text):
        with open('cache/{0}.{1}'.format(filename,type), 'w+') as file:
            file.write(text)
            file.close()

    def render_txt(self):
        # imp_word=[]
        pass

    def render_img(self):
        pass

    def generate_docx(self):
        # 1
        self.file_path = self.pro_dict.text()
        self.progressBar.setValue(1)
        # 2
        self.gen_path = self.output_dict.text()

        # 3
        if os.path.exists('cache/file.tp'):
            with open('cache/file.tp', 'r') as file:
                temp = file.read()
                file.close()
        # 4
        file_type = temp.split(';')
        # 5
        file_type_ok = True
        if len(file_type) == 0:
            QMessageBox.warning(self, '错误', '文件类型不能为空', QMessageBox.Yes)
            return
        # 6
        for item in file_type:
            if not item.startswith('.'):
                QMessageBox.warning(self, '错误', '文件类型错误', QMessageBox.Yes)
                file_type_ok=False
        # 7
        if not file_type_ok:
            return

        # 8
        if os.path.exists(self.file_path) and os.path.exists(self.gen_path):
            self.progressBar.setVisible(True)
            # 3
            self.files_list = self.visitDir(self.file_path)
            file_type.append('Makefile')
            for item in self.files_list:
                filename = os.path.splitext(item)[0]
                suffix = os.path.splitext(item)[1]
                head = filename.split('\\')[-1]
                print(head)

                if suffix in file_type:
                    data=self.read_file(item)
                    self.document.add_heading(head+suffix)
                    self.document.add_paragraph(data)
                if head == 'Makefile':
                    data = self.read_file(item)
                    self.document.add_heading(head)
                    self.document.add_paragraph(data)
            value = 0
            for i in range(8):
                time.sleep(0.05)
                value += 12.5
                self.progressBar.setValue(value)
            print(self.gen_path)
            self.document.save('{0}/{1}.docx'.format(self.gen_path, self.file_path.split('/')[-1]))
        else:
            QMessageBox.warning(self, '错误', '路径不能为空', QMessageBox.Yes)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = UserWidgets()
    win.setWindowTitle('代码文档生成器')
    win.show()
    sys.exit(app.exec_())

